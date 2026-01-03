from __future__ import annotations

import csv
import io
import json
import re
import tempfile
import zipfile
from datetime import datetime, timezone
from os import close as _close_fd
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import letter
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas

from app.storage import read_json, write_json


def _now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


def _safe_filename(name: str) -> str:
    name = re.sub(r"[^a-zA-Z0-9._-]+", "_", name).strip("._-")
    return name or "transcript"


def _srt_time(seconds: float) -> str:
    s = max(0.0, float(seconds))
    hh = int(s // 3600)
    mm = int((s % 3600) // 60)
    ss = int(s % 60)
    ms = int(round((s - int(s)) * 1000))
    if ms >= 1000:
        ss += 1
        ms -= 1000
    return f"{hh:02d}:{mm:02d}:{ss:02d},{ms:03d}"


def _vtt_time(seconds: float) -> str:
    s = max(0.0, float(seconds))
    hh = int(s // 3600)
    mm = int((s % 3600) // 60)
    ss = int(s % 60)
    ms = int(round((s - int(s)) * 1000))
    if ms >= 1000:
        ss += 1
        ms -= 1000
    return f"{hh:02d}:{mm:02d}:{ss:02d}.{ms:03d}"


def _normalize_segments(transcript: dict, edited_segments: list[dict[str, Any]]) -> list[dict[str, Any]]:
    if edited_segments:
        return [
            {
                "start": float(s["start"]),
                "end": float(s["end"]),
                "text": str(s.get("text", "")),
                "speaker": str(s["speaker"]) if s.get("speaker") else None,
            }
            for s in edited_segments
        ]

    speaker_segs = (transcript or {}).get("speaker_segments")
    if isinstance(speaker_segs, list) and speaker_segs:
        normalized: list[dict[str, Any]] = []
        for s in speaker_segs:
            if not isinstance(s, dict):
                continue
            normalized.append(
                {
                    "start": float(s.get("start", 0.0)),
                    "end": float(s.get("end", 0.0)),
                    "text": str(s.get("text", "")),
                    "speaker": str(s.get("speaker")) if s.get("speaker") else None,
                }
            )
        if normalized:
            return normalized

    ts = (transcript or {}).get("timestamps") or {}
    segs = ts.get("segment") or []
    normalized: list[dict[str, Any]] = []
    for s in segs:
        if not isinstance(s, dict):
            continue
        normalized.append(
            {
                "start": float(s.get("start", 0.0)),
                "end": float(s.get("end", 0.0)),
                "text": str(s.get("segment") if "segment" in s else s.get("text", "")),
                "speaker": None,
            }
        )
    if normalized:
        return normalized

    text = str((transcript or {}).get("text", "")).strip()
    if not text:
        return []
    return [{"start": 0.0, "end": 0.0, "text": text, "speaker": None}]


def _build_srt(segments: list[dict[str, Any]]) -> str:
    lines: list[str] = []
    for i, s in enumerate(segments, start=1):
        lines.append(str(i))
        lines.append(f"{_srt_time(s['start'])} --> {_srt_time(s['end'])}")
        speaker = str(s.get("speaker") or "").strip()
        text = str(s.get("text", "")).strip()
        if speaker:
            text = f"{speaker}: {text}".strip()
        lines.append(text)
        lines.append("")
    return "\n".join(lines).strip() + "\n"


def _build_vtt(segments: list[dict[str, Any]]) -> str:
    lines: list[str] = ["WEBVTT", ""]
    for s in segments:
        lines.append(f"{_vtt_time(s['start'])} --> {_vtt_time(s['end'])}")
        speaker = str(s.get("speaker") or "").strip()
        text = str(s.get("text", "")).strip()
        if speaker:
            text = f"{speaker}: {text}".strip()
        lines.append(text)
        lines.append("")
    return "\n".join(lines).strip() + "\n"


def _build_csv(segments: list[dict[str, Any]]) -> str:
    out = io.StringIO()
    writer = csv.writer(out)
    writer.writerow(["start_seconds", "end_seconds", "speaker", "text"])
    for s in segments:
        writer.writerow(
            [
                f"{float(s['start']):.3f}",
                f"{float(s['end']):.3f}",
                str(s.get("speaker") or ""),
                str(s.get("text", "")),
            ]
        )
    return out.getvalue()


def _build_markdown(transcript: dict, edited_text: str, segments: list[dict[str, Any]]) -> str:
    text = (edited_text or transcript.get("text") or "").strip()
    lines = ["# Transcript", "", text, ""]
    if segments:
        lines.extend(["## Segments", ""])
        for s in segments:
            spk = str(s.get("speaker") or "").strip()
            prefix = f"**{spk}** " if spk else ""
            lines.append(
                f"- `{_vtt_time(s['start'])} → {_vtt_time(s['end'])}` {prefix}{str(s.get('text','')).strip()}"
            )
    lines.append("")
    return "\n".join(lines)


def _build_html(transcript: dict, edited_text: str, segments: list[dict[str, Any]]) -> str:
    text = (edited_text or transcript.get("text") or "").strip()
    esc = (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )
    seg_html = ""
    if segments:
        rows = []
        for s in segments:
            spk = str(s.get("speaker") or "").strip()
            rows.append(
                f"<tr><td>{_vtt_time(s['start'])}</td><td>{_vtt_time(s['end'])}</td><td>{spk}</td><td>{str(s.get('text','')).strip()}</td></tr>"
            )
        seg_html = (
            "<h2>Segments</h2><table><thead><tr><th>Start</th><th>End</th><th>Speaker</th><th>Text</th></tr></thead><tbody>"
            + "".join(rows)
            + "</tbody></table>"
        )

    return (
        "<!doctype html><html><head><meta charset='utf-8'/>"
        "<meta name='viewport' content='width=device-width,initial-scale=1'/>"
        "<title>Transcript</title>"
        "<style>body{font-family:system-ui,Segoe UI,Roboto,Helvetica,Arial,sans-serif;max-width:900px;margin:24px auto;padding:0 16px}"
        "pre{white-space:pre-wrap} table{border-collapse:collapse;width:100%} td,th{border:1px solid #ddd;padding:6px 8px;font-size:14px}</style>"
        "</head><body>"
        "<h1>Transcript</h1>"
        f"<pre>{esc}</pre>"
        f"{seg_html}"
        "</body></html>"
    )


def _build_docx(transcript: dict, edited_text: str, segments: list[dict[str, Any]]) -> bytes:
    doc = Document()
    doc.add_heading("Transcript", level=1)
    text = (edited_text or transcript.get("text") or "").strip()
    p = doc.add_paragraph(text)
    p.style.font.size = Pt(11)

    if segments:
        doc.add_heading("Segments", level=2)
        table = doc.add_table(rows=1, cols=4)
        hdr = table.rows[0].cells
        hdr[0].text = "Start"
        hdr[1].text = "End"
        hdr[2].text = "Speaker"
        hdr[3].text = "Text"
        for s in segments:
            row = table.add_row().cells
            row[0].text = _vtt_time(s["start"])
            row[1].text = _vtt_time(s["end"])
            row[2].text = str(s.get("speaker") or "")
            row[3].text = str(s.get("text", "")).strip()

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def _build_pdf(transcript: dict, edited_text: str) -> bytes:
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)

    # Use a builtin font; attempt to register DejaVuSans if available for better unicode.
    font_name = "Helvetica"
    try:
        import os

        candidates = [
            "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
            "/Library/Fonts/DejaVuSans.ttf",
        ]
        ttf = next((p for p in candidates if os.path.exists(p)), None)
        if ttf:
            pdfmetrics.registerFont(TTFont("DejaVuSans", ttf))
            font_name = "DejaVuSans"
    except Exception:
        pass

    c.setFont(font_name, 16)
    c.drawString(72, 750, "Transcript")

    c.setFont(font_name, 11)
    text_obj = c.beginText(72, 725)
    text = (edited_text or transcript.get("text") or "").strip()
    for line in text.splitlines() or [""]:
        text_obj.textLine(line)
    c.drawText(text_obj)
    c.showPage()
    c.save()
    return buf.getvalue()


def build_export_bytes(
    job_id: str,
    transcript_path: Path,
    format: str,
    edited_text: str,
    edited_segments: list[dict[str, Any]],
) -> tuple[bytes, str, str]:
    transcript = read_json(transcript_path)
    segments = _normalize_segments(transcript, edited_segments)

    stem = _safe_filename(f"transcript_{job_id}")

    fmt = format.lower().lstrip(".")
    if fmt == "srt":
        return _build_srt(segments).encode("utf-8"), "text/plain; charset=utf-8", f"{stem}.srt"
    if fmt == "vtt":
        return _build_vtt(segments).encode("utf-8"), "text/vtt; charset=utf-8", f"{stem}.vtt"
    if fmt == "csv":
        return _build_csv(segments).encode("utf-8"), "text/csv; charset=utf-8", f"{stem}.csv"
    if fmt == "md":
        return (
            _build_markdown(transcript, edited_text, segments).encode("utf-8"),
            "text/markdown; charset=utf-8",
            f"{stem}.md",
        )
    if fmt == "html":
        return _build_html(transcript, edited_text, segments).encode("utf-8"), "text/html; charset=utf-8", f"{stem}.html"
    if fmt == "docx":
        return _build_docx(transcript, edited_text, segments), "application/vnd.openxmlformats-officedocument.wordprocessingml.document", f"{stem}.docx"
    if fmt == "pdf":
        return _build_pdf(transcript, edited_text), "application/pdf", f"{stem}.pdf"
    if fmt == "dote":
        payload = {
            "format": "dote",
            "created_at": _now_iso(),
            "job_id": job_id,
            "model_name": transcript.get("model_name"),
            "device": transcript.get("device"),
            "text": (edited_text or transcript.get("text") or ""),
            "segments": segments,
            "timestamps": transcript.get("timestamps"),
            "diarization": transcript.get("diarization"),
        }
        return json.dumps(payload, ensure_ascii=False, indent=2).encode("utf-8"), "application/json; charset=utf-8", f"{stem}.dote"

    raise ValueError(f"Unsupported export format: {format}")


def build_whisper_bundle(
    job_id: str,
    original_audio_path: Path,
    wav_audio_path: Path,
    transcript_path: Path,
    edited_text: str,
    edited_segments: list[dict[str, Any]],
    out_dir: Path,
) -> Path:
    transcript = read_json(transcript_path)
    segments = _normalize_segments(transcript, edited_segments)

    edits = {
        "created_at": _now_iso(),
        "job_id": job_id,
        "edited_text": edited_text,
        "edited_segments": segments,
    }

    write_json(out_dir / "edits.json", edits)

    fd, tmp_path = tempfile.mkstemp(prefix=f"bundle_{job_id}_", suffix=".whisper", dir=str(out_dir))
    _close_fd(fd)

    bundle_path = Path(tmp_path)
    with zipfile.ZipFile(bundle_path, "w", compression=zipfile.ZIP_DEFLATED) as z:
        if original_audio_path.exists():
            z.write(original_audio_path, arcname=original_audio_path.name)
        if wav_audio_path.exists():
            z.write(wav_audio_path, arcname=wav_audio_path.name)
        z.write(transcript_path, arcname="transcript.json")

        edits_bytes = json.dumps(edits, ensure_ascii=False, indent=2).encode("utf-8")
        z.writestr("edits.json", edits_bytes)

        meta = {
            "format": "whisper-bundle",
            "created_at": _now_iso(),
            "job_id": job_id,
            "model_name": transcript.get("model_name"),
            "device": transcript.get("device"),
        }
        z.writestr("meta.json", json.dumps(meta, ensure_ascii=False, indent=2).encode("utf-8"))

    return bundle_path
